#!/usr/bin/env python3
"""
Script to update the Alfaisal University HR Job Description template
for Postdoctoral Research Fellow in Cognitive Robotics at AlfaisalX
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# Path to the template and output
TEMPLATE_PATH = "/Applications/XAMPP/xamppfiles/htdocs/aniskoubaa.org/alfaisalx/Postdoctoral Research Fellow - 15 October 2025.doc"
OUTPUT_PATH = "/Applications/XAMPP/xamppfiles/htdocs/aniskoubaa.org/alfaisalx/Postdoctoral Research Fellow - Cognitive Robotics - AlfaisalX.docx"

# New content for Cognitive Robotics position
JOB_DETAILS = {
    "JOB TITLE:": "Postdoctoral Research Fellow - Cognitive Robotics",
    "DEPARTMENT:": "AlfaisalX - Center of Excellence for Cognitive Robotics & Autonomous Agents",
    "DIVISION:": "COE&AC",
    "JOB FAMILY:": "Faculty- Academic",
    "LOCATION:": "Alfaisal University",
    "REPORTS TO:": "Director of AlfaisalX (Prof. Anis Koubaa)",
    "GRADE:": "07",
}

JOB_SUMMARY = """The Postdoctoral Research Fellow at AlfaisalX supports strategic research initiatives focused on developing next-generation cognitive robotics and autonomous systems. This includes contributing to interdisciplinary research in areas such as robot perception, autonomous navigation, multi-agent systems, UAV autonomy, and agentic AI workflows, with applications spanning smart cities, healthcare, logistics, and industrial automation. The role advances the center's mission by integrating artificial intelligence, machine learning, robot control, and cloud robotics to address real-world challenges through innovative, translational solutions aligned with Saudi Vision 2030."""

PRIMARY_DUTIES = [
    "Produces high-quality publications and delivers presentations at national or international conferences to disseminate findings in cognitive robotics, autonomous systems, and AI.",
    "Actively contributes to interdisciplinary research initiatives that align with AlfaisalX's strategic goals in robotics, UAVs, and agentic AI, fostering collaboration across fields to address complex challenges.",
    "Collaborates with academic, industrial, and governmental partners to facilitate translational and applied research that bridges fundamental robotics science with practical engineering applications.",
    "Facilitates interdisciplinary communication by organizing seminars, workshops, or discussion forums that promote knowledge exchange in robotics, autonomous systems, and advanced AI fields.",
    "Participates in the preparation of research proposals for external funding and supports reporting to ensure sustained project resources and compliance with grant requirements.",
    "Designs, conducts, and analyzes experimental or computational research in cognitive robotics, autonomous navigation, and multi-agent coordination to generate valid, reproducible results.",
    "Supervises and mentors undergraduate and graduate students to support their research development, skill acquisition, and contribution to AlfaisalX goals.",
    "Develops and prototypes robotic systems, autonomous platforms, or AI-driven controllers, ensuring usability, safety, and innovation to advance experimental capabilities.",
    "Builds and refines computational models, simulation environments, or digital twins that support data-driven robotics investigations and knowledge advancement.",
    "Contributes to the preparation and dissemination of research findings through publications, technical reports, and presentations to scientific and professional audiences.",
    "Collaborates with team members to meet project milestones, promote interdisciplinary knowledge exchange in robotics and AI, and ensure timely progress.",
    "Supports the development of algorithms, control systems, or machine learning models applicable to ongoing robotics and autonomous systems projects.",
    "Maintains research documentation, datasets, and code in a professional, organized manner to ensure accuracy, reproducibility, and compliance with data management standards.",
    "Performs all duties in a professional, effective, and confidential manner.",
    "Performs all other related duties as required or assigned by the Director of AlfaisalX.",
]

EDUCATION_REQUIREMENTS = """Doctor of Philosophy (Ph.D.) in Computer Science, Robotics Engineering, Artificial Intelligence, Electrical & Computer Engineering, Mechatronics, or an equivalent degree from an accredited institution with focus on autonomous systems, cognitive robotics, or related fields."""

PROFESSIONAL_EXPERIENCE = """Minimum of one year of research experience supporting or conducting projects involving robotics, autonomous systems, machine learning, computer vision, or related computational and experimental work in AI and robotics fields."""

BEHAVIORAL_COMPETENCIES = [
    ("Confidentiality and Discretion", "Advanced"),
    ("Work Ethics and Value", "Advanced"),
    ("Interactive Communication", "Advanced"),
    ("Creativity and Innovation", "Advanced"),
    ("Research Initiative and Independence", "Competent"),
]

FUNCTIONAL_COMPETENCIES = [
    ("Develops and implements robotics research protocols following ethical standards", "Advanced"),
    ("Demonstrated with robotic platforms, ROS/ROS2, and AI/ML frameworks", "Advanced"),
    ("Proficiency in scientific writing and presenting robotics research findings", "Advanced"),
    ("Computer literate with proficiency in Python, C++, and simulation tools", "Advanced"),
    ("Fluency in oral and written English language", "Advanced"),
]

CORE_COMPETENCIES = [
    ("Collect and Interpret Robotics Data", "Competent"),
    ("Algorithm Development and Evaluation Methods", "Competent"),
    ("Document and Track Experimental Robotics Data", "Competent"),
    ("Data Accuracy and Research Integrity", "Advanced"),
    ("Robotics Tools and Technology Proficiency", "Advanced"),
]


def update_document():
    """Read the template and create updated document with new content."""
    
    # Try to open the .doc file - python-docx may need the file to be .docx
    # We'll create a new document based on the template structure
    print(f"Creating new document based on template structure...")
    
    doc = Document()
    
    # Set narrow margins
    for section in doc.sections:
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
    
    # Add header with title
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run("JOB DESCRIPTION & SPECIFICATION")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "Arial"
    
    subheader = doc.add_paragraph()
    subheader.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subheader.add_run("Human Resources Department")
    run.font.size = Pt(12)
    run.font.name = "Arial"
    
    doc.add_paragraph()  # spacing
    
    # Create job details table
    job_table = doc.add_table(rows=4, cols=4)
    job_table.style = 'Table Grid'
    
    job_data = [
        [("JOB TITLE:", True), (JOB_DETAILS["JOB TITLE:"], False), ("GRADE:", True), (JOB_DETAILS["GRADE:"], False)],
        [("DEPARTMENT:", True), (JOB_DETAILS["DEPARTMENT:"], False), ("DIVISION:", True), (JOB_DETAILS["DIVISION:"], False)],
        [("JOB FAMILY:", True), (JOB_DETAILS["JOB FAMILY:"], False), ("LOCATION:", True), (JOB_DETAILS["LOCATION:"], False)],
        [("REPORTS TO:", True), (JOB_DETAILS["REPORTS TO:"], False), ("", False), ("", False)],
    ]
    
    for i, row_data in enumerate(job_data):
        row = job_table.rows[i]
        for j, (text, is_bold) in enumerate(row_data):
            cell = row.cells[j]
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run(text)
            run.bold = is_bold
            run.font.size = Pt(10)
            run.font.name = "Arial"
    
    doc.add_paragraph()  # spacing
    
    # Section 1: Job Summary
    section1 = doc.add_paragraph()
    run = section1.add_run("1. JOB SUMMARY:")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Arial"
    
    summary_table = doc.add_table(rows=1, cols=1)
    summary_table.style = 'Table Grid'
    cell = summary_table.cell(0, 0)
    para = cell.paragraphs[0]
    run = para.add_run(JOB_SUMMARY)
    run.font.size = Pt(10)
    run.font.name = "Arial"
    
    doc.add_paragraph()  # spacing
    
    # Section 2: Primary Duties & Responsibilities
    section2 = doc.add_paragraph()
    run = section2.add_run("2. PRIMARY DUTIES & RESPONSIBILITIES:")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Arial"
    
    duties_table = doc.add_table(rows=len(PRIMARY_DUTIES), cols=2)
    duties_table.style = 'Table Grid'
    
    for i, duty in enumerate(PRIMARY_DUTIES):
        row = duties_table.rows[i]
        # Number column
        cell_num = row.cells[0]
        para = cell_num.paragraphs[0]
        run = para.add_run(f"{i+1}.")
        run.font.size = Pt(10)
        run.font.name = "Arial"
        # Duty column
        cell_duty = row.cells[1]
        para = cell_duty.paragraphs[0]
        run = para.add_run(duty)
        run.font.size = Pt(10)
        run.font.name = "Arial"
    
    # Set column widths
    for row in duties_table.rows:
        row.cells[0].width = Inches(0.4)
        row.cells[1].width = Inches(6.5)
    
    doc.add_page_break()
    
    # Section 3: Qualifications and Competencies
    section3 = doc.add_paragraph()
    run = section3.add_run("3. QUALIFICATIONS AND COMPETENCIES:")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Arial"
    
    edu_header = doc.add_paragraph()
    run = edu_header.add_run("EDUCATION/LICENSES:")
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = "Arial"
    
    edu_para = doc.add_paragraph()
    run = edu_para.add_run("• " + EDUCATION_REQUIREMENTS)
    run.font.size = Pt(10)
    run.font.name = "Arial"
    
    exp_header = doc.add_paragraph()
    run = exp_header.add_run("PROFESSIONAL EXPERIENCE:")
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = "Arial"
    
    exp_para = doc.add_paragraph()
    run = exp_para.add_run("• " + PROFESSIONAL_EXPERIENCE)
    run.font.size = Pt(10)
    run.font.name = "Arial"
    
    doc.add_paragraph()  # spacing
    
    # Section 4: Competencies
    section4 = doc.add_paragraph()
    run = section4.add_run("4. COMPETENCIES:")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Arial"
    
    # Helper function to create competency table
    def add_competency_table(title, competencies):
        # Title row
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run(title)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = "Arial"
        
        # Create table with header
        table = doc.add_table(rows=len(competencies) + 1, cols=5)
        table.style = 'Table Grid'
        
        # Header row
        header_row = table.rows[0]
        headers = ["", "", "Basic", "Competent", "Advanced"]
        header_row.cells[0].merge(header_row.cells[1])
        
        para = header_row.cells[0].paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run("")
        
        # Proficiency Level header spanning 3 columns
        for j, header_text in enumerate(["", "", "Basic", "Competent", "Advanced"]):
            if j >= 2:
                para = header_row.cells[j].paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run(header_text)
                run.bold = True
                run.font.size = Pt(9)
                run.font.name = "Arial"
        
        # Data rows
        for i, (comp_name, level) in enumerate(competencies):
            row = table.rows[i + 1]
            # Number
            para = row.cells[0].paragraphs[0]
            run = para.add_run(f"{i+1}.")
            run.font.size = Pt(9)
            run.font.name = "Arial"
            # Competency name
            para = row.cells[1].paragraphs[0]
            run = para.add_run(comp_name)
            run.font.size = Pt(9)
            run.font.name = "Arial"
            # Mark the level
            level_map = {"Basic": 2, "Competent": 3, "Advanced": 4}
            if level in level_map:
                para = row.cells[level_map[level]].paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run("X")
                run.bold = True
                run.font.size = Pt(9)
                run.font.name = "Arial"
        
        doc.add_paragraph()  # spacing
    
    add_competency_table("BEHAVIORAL COMPETENCIES", BEHAVIORAL_COMPETENCIES)
    add_competency_table("FUNCTIONAL COMPETENCIES", FUNCTIONAL_COMPETENCIES)
    add_competency_table("CORE COMPETENCIES", CORE_COMPETENCIES)
    
    # Approval Section
    approval_header = doc.add_paragraph()
    approval_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = approval_header.add_run("APPROVAL")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Arial"
    
    approval_table = doc.add_table(rows=4, cols=3)
    approval_table.style = 'Table Grid'
    
    # Header row
    header_row = approval_table.rows[0]
    for j, text in enumerate(["", "SIGNATURE", "DATE"]):
        para = header_row.cells[j].paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(text)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = "Arial"
    
    approval_data = [
        "Prepared by: Department Head",
        "Reviewed by: HR Planning & Development Manager",
        "Approved by: Director, Human Resources",
    ]
    
    for i, text in enumerate(approval_data):
        row = approval_table.rows[i + 1]
        para = row.cells[0].paragraphs[0]
        run = para.add_run(text)
        run.italic = True
        run.font.size = Pt(10)
        run.font.name = "Arial"
    
    doc.add_paragraph()  # spacing
    
    # Disclaimer
    disclaimer_header = doc.add_paragraph()
    disclaimer_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = disclaimer_header.add_run("DISCLAIMER")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Arial"
    
    disclaimer_table = doc.add_table(rows=1, cols=1)
    disclaimer_table.style = 'Table Grid'
    cell = disclaimer_table.cell(0, 0)
    para = cell.paragraphs[0]
    run = para.add_run("The above statement describes the general nature and level of work performed by individuals assigned to this position. This is not intended to be an exhaustive list of all responsibilities and duties required of personnel as classified.")
    run.font.size = Pt(10)
    run.font.name = "Arial"
    
    doc.add_paragraph()  # spacing
    
    # Signature section
    sig_table = doc.add_table(rows=3, cols=3)
    sig_table.style = 'Table Grid'
    
    # Header row
    header_row = sig_table.rows[0]
    for j, text in enumerate(["", "SIGNATURE", "DATE"]):
        para = header_row.cells[j].paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(text)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = "Arial"
    
    sig_data = ["Employee's Name:", "Direct Superior's Name:"]
    for i, text in enumerate(sig_data):
        row = sig_table.rows[i + 1]
        para = row.cells[0].paragraphs[0]
        run = para.add_run(text)
        run.italic = True
        run.font.size = Pt(10)
        run.font.name = "Arial"
    
    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    run = footer.add_run("Updated: 15 October 2025")
    run.font.size = Pt(8)
    run.font.name = "Arial"
    footer.add_run("\t\t\t\t\t")
    run2 = footer.add_run("Prepared for AlfaisalX")
    run2.font.size = Pt(8)
    run2.font.name = "Arial"
    
    # Save the document
    doc.save(OUTPUT_PATH)
    print(f"Document saved to: {OUTPUT_PATH}")
    print("Done!")
    

if __name__ == "__main__":
    update_document()
